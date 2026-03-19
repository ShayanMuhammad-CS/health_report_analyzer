import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(0, 51, 102)

def generate_documentation():
    doc = Document()

    # Document Title
    title = doc.add_heading('DOCUS AI - Comprehensive Project Documentation & Demonstration Guide', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('A Complete Guide from Scratch to Finish').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # Introduction
    add_heading(doc, '1. Introduction: What is DOCUS AI?', 1)
    doc.add_paragraph(
        "DOCUS AI (formerly known as HIA - Health Report Analyzer) is an advanced, AI-powered "
        "medical application designed to demystify complex blood reports. It empowers users "
        "by extracting text from uploaded health reports, analyzing the data using state-of-the-art "
        "Large Language Models (Groq Llama 3.3 70B), and providing an interactive 'chat' interface "
        "for follow-up questions."
    )
    doc.add_paragraph(
        "Built with a focus on an ultra-premium user experience, the system features a stunning "
        "glassmorphism UI, interactive elements, custom typography, and secure user authentication."
    )

    # Core Technologies
    add_heading(doc, '2. Core Technology Stack', 1)
    tech_stack = [
        "Frontend: Streamlit (heavily customized with pure CSS for a modern, animated, glassmorphic design).",
        "Authentication & Database: Supabase (manages user signups, logins, and chat history).",
        "AI Engine: Groq API utilizing the Llama-3.3-70b-versatile model for lightning-fast inference.",
        "RAG Pipeline: LangChain, HuggingFace Embeddings (all-MiniLM-L6-v2), and FAISS Vector Database for contextual follow-up Q&A.",
        "PDF Processing: pdfplumber for robust text extraction and rule-based validation to ensure genuine reports."
    ]
    for tech in tech_stack:
        doc.add_paragraph(tech, style='List Bullet')

    # Step-by-Step Flow
    add_heading(doc, '3. The Complete User Journey (From Scratch to Finish)', 1)
    
    add_heading(doc, 'Step 1: Welcome & Authentication', 2)
    doc.add_paragraph(
        "Users arrive at a beautifully animated login page. The application natively integrates with "
        "Supabase. New users can sign up using an email and password, while returning users securely log in. "
        "The session is managed seamlessly behind the scenes."
    )

    add_heading(doc, 'Step 2: Dashboard & Session Creation', 2)
    doc.add_paragraph(
        "Once logged in, users are greeted with their name in the header. The sidebar, designed with a sleek "
        "glass-like transparency, allows them to create a 'New Session' or view past chat histories. "
        "Every session acts as an independent medical consultation workspace."
    )

    add_heading(doc, 'Step 3: Uploading the Blood Report', 2)
    doc.add_paragraph(
        "In a new session, users provide basic demographic info (Age, Gender, Medical History) and "
        "can either upload their actual PDF blood report (max 20 pages) or try a pre-loaded sample. "
        "The system instantly reads the PDF using 'pdfplumber' and validates that it contains relevant "
        "medical terminology (like hemoglobin, glucose, etc.)."
    )

    add_heading(doc, 'Step 4: AI Analysis (The Magic Happens)', 2)
    doc.add_paragraph(
        "Once verified, the 'AnalysisAgent' takes over. It verifies rate limits (max 15 analyses per user/day). "
        "It then compiles the patient's data, age, gender, and the extracted report text. The system leverages "
        "'In-Context Learning', pulling similar past analyses from its knowledge base to enhance accuracy, "
        "and prompts the Groq LLM to generate a highly detailed, medically sound, and structured analysis of the report. "
        "This result is dynamically rendered in a premium summary box on the UI."
    )

    add_heading(doc, 'Step 5: Interactive Chat (RAG Capability)', 2)
    doc.add_paragraph(
        "After the initial analysis is presented, an iMessage-style chat interface appears below. Here, "
        "users can ask follow-up questions (e.g., 'What does my high cholesterol mean for my diet?'). "
        "Behind the scenes, the 'ChatAgent' splits the blood report into chunks, converts them into numeric vectors "
        "using HuggingFace models, and stores them in a FAISS database. It then 'contextualizes' the user's question "
        "based on chat history, searches the FAISS database for relevant report sections, and feeds "
        "that exact context to the AI to answer specifically and accurately. It's a true RAG (Retrieval-Augmented Generation) pipeline."
    )

    # First-Time Demonstration Guide
    doc.add_page_break()
    add_heading(doc, '4. The Ultimate Pitch: First-Time Demonstration Guide', 1)
    doc.add_paragraph(
        "To impress everyone during a live demonstration, follow this exact script and flow to showcase the system's power. "
        "Make sure to highlight both the technical brilliance and the unparalleled UI/UX."
    )

    demo_steps = [
        "1. Start at the Login Screen: Point out the customized aurora gradient background and the seamless Supabase integration. Mention that security and aesthetics go hand in hand.",
        "2. Create a Session: Click 'New Session'. Emphasize the speed and the modern glassmorphism sidebar. The UI doesn't look like standard Streamlit; it looks like a premium SaaS product.",
        "3. Use the Sample File: Select the 'Try Sample Blood Report' option. Fill in dummy data (e.g., Age 45, Male, History of Hypertension). Show how the premium styled button reacts smoothly to clicks.",
        "4. The Initial Analysis Box: Click 'Analyze Report'. While the skeleton loader spins, explain that the Groq LLM (Llama 3.3) is processing the data at lightning speeds while the system enforces rate limits to prevent abuse.",
        "5. Dive into the RAG Chat: Once the analysis appears, scroll down to the chat. Type a specific question like: 'Based on my HDL, what specific foods should I avoid?'.",
        "6. Explain RAG Live: As it answers, explain that it's not just using general AI knowledge; it is retrieving exact chunks of the uploaded PDF using vector embeddings (FAISS) to answer accurately based on the user's specific report.",
        "7. Final Polish: Point out the beautiful chat bubbles (blue for user, dark glass for AI) and the floating chat input pill. It's a flawless finish."
    ]
    for step in demo_steps:
        doc.add_paragraph(step)

    # Technical Milestones
    add_heading(doc, '5. Standout Technical Features (For Developers/Judges)', 1)
    features = [
        "Complete UI Overhaul: Streaming custom CSS directly into Streamlit to hide default branding, add 15s gradient animations, and implement iMessage-style flex-row-reverse chat bubbles.",
        "In-Context Learning Knowledge Base: The system remembers previous successful analyses for specific symptoms (like high glucose in a 50yo male) and feeds them back into the prompt for extreme accuracy over time.",
        "Rate Limiting & Defensive Design: Built-in counters (15 per 24 hours) ensure API costs remain perfectly managed, using clever datetime delta checks.",
        "Session Based Chat Histories: Supabase integration securely stores and fetches chat histories tied perfectly to independent session UUIDs."
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')

    # Conclusion
    add_heading(doc, '6. Conclusion', 1)
    doc.add_paragraph(
        "DOCUS AI is more than just a wrapper; it is a full-stack, enterprise-ready application that solves "
        "a real-world problem (understanding medical data) natively, securely, and beautifully."
    )

    doc.save('DOCUS_AI_Comprehensive_Explanation.docx')
    print("Document successfully generated: DOCUS_AI_Comprehensive_Explanation.docx")

if __name__ == '__main__':
    generate_documentation()
